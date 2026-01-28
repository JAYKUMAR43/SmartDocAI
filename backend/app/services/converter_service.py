import os
import sys
import uuid
import shutil
import subprocess
import platform
from PIL import Image
from app.core import config
from pdf2docx import Converter as PDF2DocxConverter
import fitz # PyMuPDF
from app.services import pdf_service, ai_service

# Platform detection
IS_WINDOWS = platform.system() == "Windows"

# Conditional imports for Windows-only libraries
if IS_WINDOWS:
    try:
        import comtypes.client
        from docx2pdf import convert as docx2pdf_convert
        WINDOWS_COM_AVAILABLE = True
    except ImportError:
        WINDOWS_COM_AVAILABLE = False
else:
    WINDOWS_COM_AVAILABLE = False


def pdf_to_docx(input_path: str, output_path: str):
    """Converts PDF to DOCX using pdf2docx."""
    try:
        cv = PDF2DocxConverter(input_path)
        cv.convert(output_path, start=0, end=None)
        cv.close()
        return True
    except Exception as e:
        raise Exception(f"PDF to DOCX failed: {str(e)}")

def docx_to_pdf_libreoffice(input_path: str, output_path: str):
    """Converts DOCX to PDF using LibreOffice (cross-platform)."""
    try:
        # LibreOffice headless conversion
        output_dir = os.path.dirname(output_path)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, input_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
        
        # LibreOffice creates file with same name but .pdf extension
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
        
        if os.path.exists(generated_pdf) and generated_pdf != output_path:
            shutil.move(generated_pdf, output_path)
        
        return True
    except Exception as e:
        raise Exception(f"DOCX to PDF failed: {str(e)}")


def docx_to_pdf(input_path: str, output_path: str):
    """Converts DOCX to PDF (platform-aware)."""
    if WINDOWS_COM_AVAILABLE:
        try:
            docx2pdf_convert(input_path, output_path)
            return True
        except Exception as e:
            raise Exception(f"DOCX to PDF failed: {str(e)}")
    else:
        return docx_to_pdf_libreoffice(input_path, output_path)

def excel_to_pdf_libreoffice(input_path: str, output_path: str):
    """Converts Excel to PDF using LibreOffice (cross-platform)."""
    try:
        output_dir = os.path.dirname(output_path)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, input_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
        
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
        
        if os.path.exists(generated_pdf) and generated_pdf != output_path:
            shutil.move(generated_pdf, output_path)
        
        return True
    except Exception as e:
        raise Exception(f"Excel to PDF failed: {str(e)}")


def excel_to_pdf(input_path: str, output_path: str):
    """Converts Excel to PDF (platform-aware)."""
    if WINDOWS_COM_AVAILABLE:
        try:
            excel = comtypes.client.CreateObject("Excel.Application")
            excel.Visible = False
            wb = excel.Workbooks.Open(os.path.abspath(input_path))
            # 0 = xlTypePDF
            wb.ExportAsFixedFormat(0, os.path.abspath(output_path))
            wb.Close()
            excel.Quit()
            return True
        except Exception as e:
            raise Exception(f"Excel to PDF failed: {str(e)}")
    else:
        return excel_to_pdf_libreoffice(input_path, output_path)

def ppt_to_pdf_libreoffice(input_path: str, output_path: str):
    """Converts PowerPoint to PDF using LibreOffice (cross-platform)."""
    try:
        output_dir = os.path.dirname(output_path)
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, input_path],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")
        
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
        
        if os.path.exists(generated_pdf) and generated_pdf != output_path:
            shutil.move(generated_pdf, output_path)
        
        return True
    except Exception as e:
        raise Exception(f"PPT to PDF failed: {str(e)}")


def ppt_to_pdf(input_path: str, output_path: str):
    """Converts PowerPoint to PDF (platform-aware)."""
    if WINDOWS_COM_AVAILABLE:
        try:
            powerpoint = comtypes.client.CreateObject("Powerpoint.Application")
            # 32 = ppFixedFormatTypePDF
            pres = powerpoint.Presentations.Open(os.path.abspath(input_path), WithWindow=False)
            pres.ExportAsFixedFormat(os.path.abspath(output_path), 32)
            pres.Close()
            powerpoint.Quit()
            return True
        except Exception as e:
            raise Exception(f"PPT to PDF failed: {str(e)}")
    else:
        return ppt_to_pdf_libreoffice(input_path, output_path)

def ppt_to_images_libreoffice(input_path: str, output_dir: str, format: str = "PNG"):
    """Converts PowerPoint slides to images using LibreOffice + PyMuPDF."""
    try:
        # First convert PPT to PDF using LibreOffice
        temp_pdf = os.path.join(output_dir, f"temp_{uuid.uuid4()}.pdf")
        ppt_to_pdf_libreoffice(input_path, temp_pdf)
        
        # Then convert PDF pages to images using PyMuPDF
        doc = fitz.open(temp_pdf)
        generated_files = []
        
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=150)
            img_path = os.path.join(output_dir, f"slide_{i+1}.{format.lower()}")
            
            if format.upper() == "PNG":
                pix.save(img_path)
            else:
                # Convert to JPG
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img.save(img_path, "JPEG")
            
            generated_files.append(img_path)
        
        doc.close()
        
        # Clean up temp PDF
        if os.path.exists(temp_pdf):
            os.remove(temp_pdf)
        
        return generated_files
    except Exception as e:
        raise Exception(f"PPT to Images failed: {str(e)}")


def ppt_to_images(input_path: str, output_dir: str, format: str = "PNG"):
    """Converts PowerPoint slides to images (platform-aware)."""
    if WINDOWS_COM_AVAILABLE:
        try:
            powerpoint = comtypes.client.CreateObject("Powerpoint.Application")
            pres = powerpoint.Presentations.Open(os.path.abspath(input_path), WithWindow=False)
            
            generated_files = []
            for i, slide in enumerate(pres.Slides):
                img_path = os.path.join(output_dir, f"slide_{i+1}.{format.lower()}")
                # 17 = ppSaveAsPNG, 18 = ppSaveAsJPG
                save_type = 17 if format.upper() == "PNG" else 18
                slide.Export(os.path.abspath(img_path), format.upper())
                generated_files.append(img_path)
                
            pres.Close()
            powerpoint.Quit()
            return generated_files
        except Exception as e:
            raise Exception(f"PPT to Images failed: {str(e)}")
    else:
        return ppt_to_images_libreoffice(input_path, output_dir, format)

def images_to_pdf(image_paths: list[str], output_path: str):
    """Converts a list of images to a single PDF."""
    return pdf_service.create_pdf_from_images(image_paths, output_path)

def text_to_pdf(input_path: str, output_path: str):
    """Converts text file to PDF."""
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return pdf_service.create_pdf_from_text(text, output_path)
    except Exception as e:
        raise Exception(f"Text to PDF failed: {str(e)}")

def markdown_to_docx(markdown_content: str, output_path: str):
    """Converts markdown text to a formatted DOCX file."""
    import docx
    doc = docx.Document()
    for line in markdown_content.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    doc.save(output_path)
    return True

def create_document_from_content(content: str, format: str, output_path: str):
    """Creates a document (DOCX/PDF) from text content."""
    if format.lower() == 'docx':
        return markdown_to_docx(content, output_path)
    elif format.lower() == 'pdf':
        # Create DOCX first then convert to PDF for better formatting than simple text
        temp_docx = output_path.replace(".pdf", ".docx")
        markdown_to_docx(content, temp_docx)
        try:
            docx_to_pdf(temp_docx, output_path)
        finally:
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
        return True
    else:
        raise ValueError(f"Unsupported format: {format}")

def ocr_and_rebuild(input_path: str, target_format: str, output_path: str):
    """
    Uses AI to perform OCR and rebuild the document in the target format.
    """
    
    # 1. Get AI Reconstruction (Markdown)
    markdown_content = ai_service.recreate_document_from_image(input_path)
    
    # 2. Convert Markdown to Target Format
    if target_format.lower() == 'docx':
        markdown_to_docx(markdown_content, output_path)
    elif target_format.lower() == 'pdf':
        create_document_from_content(markdown_content, 'pdf', output_path)
    else:
        # Fallback to plain text
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
    
    return True

def convert_document(input_path: str, target_format: str, session_id: str = None) -> str:
    """
    Main entry point for universal conversions.
    """
    file_ext = os.path.splitext(input_path)[1].lower()
    
    # Define output session directory
    if session_id:
        output_dir = os.path.join(config.TEMP_DIR, session_id)
    else:
        session_id = str(uuid.uuid4())
        output_dir = os.path.join(config.TEMP_DIR, session_id)
    
    os.makedirs(output_dir, exist_ok=True)
    
    output_filename = f"converted_{uuid.uuid4()}.{target_format.lower()}"
    output_path = os.path.join(output_dir, output_filename)

    # Scanned Document Detection
    is_image = file_ext in ['.jpg', '.jpeg', '.png', '.webp', '.bmp']
    is_scanned_pdf = False
    if file_ext == '.pdf':
        try:
            text = pdf_service.extract_text_from_pdf(input_path)
            if len(text.strip()) < 10:
                is_scanned_pdf = True
        except: pass

    if (is_image or is_scanned_pdf) and target_format.lower() in ['docx', 'pdf', 'xlsx']:
        ocr_and_rebuild(input_path, target_format, output_path)
        return output_path

    # Standard Conversions
    try:
        # PDF -> X
        if file_ext == '.pdf':
            if target_format.lower() == 'docx': pdf_to_docx(input_path, output_path)
            elif target_format.lower() in ['jpg', 'png']:
                doc = fitz.open(input_path)
                pix = doc[0].get_pixmap() # Page 1
                pix.save(output_path)
                doc.close()
            elif target_format.lower() == 'txt':
                text = pdf_service.extract_text_from_pdf(input_path)
                with open(output_path, "w", encoding="utf-8") as f: f.write(text)
            else: raise Exception(f"PDF to {target_format} not directly supported.")

        # WORD -> X
        elif file_ext in ['.docx', '.doc']:
            if target_format.lower() == 'pdf': docx_to_pdf(input_path, output_path)
            elif target_format.lower() == 'txt':
                import docx
                doc = docx.Document(input_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                with open(output_path, "w", encoding="utf-8") as f: f.write(text)
            else: raise Exception(f"Word to {target_format} not supported.")

        # EXCEL -> X
        elif file_ext in ['.xlsx', '.xls']:
            if target_format.lower() == 'pdf': excel_to_pdf(input_path, output_path)
            else: raise Exception(f"Excel to {target_format} not supported.")

        # PPT -> X
        elif file_ext in ['.pptx', '.ppt']:
            if target_format.lower() == 'pdf': ppt_to_pdf(input_path, output_path)
            elif target_format.lower() in ['jpg', 'png']:
                # Note: This might return multiple files, we take the first for consistency here
                # In batch mode we handle multifile differently
                ppt_to_images(input_path, output_dir, format=target_format)
                # output_path is already set, but ppt_to_images writes slide_1.png etc.
                # Let's fix output_path to point to the first slide
                output_path = os.path.join(output_dir, f"slide_1.{target_format.lower()}")
            else: raise Exception(f"PPT to {target_format} not supported.")

        # TXT -> X
        elif file_ext == '.txt':
            if target_format.lower() == 'pdf': text_to_pdf(input_path, output_path)
            else: shutil.copy(input_path, output_path)

        # IMAGES -> X
        elif is_image:
            if target_format.lower() == 'pdf': images_to_pdf([input_path], output_path)
            else: shutil.copy(input_path, output_path)

        else:
            raise Exception(f"Unsupported input format: {file_ext}")

    except Exception as e:
        # Intermediate conversion logic: X -> PDF -> Y
        if target_format.lower() == 'docx' and file_ext in ['.pptx', '.ppt', '.xlsx', '.xls']:
            temp_pdf = os.path.join(output_dir, f"temp_{uuid.uuid4()}.pdf")
            convert_document(input_path, 'pdf', session_id)
            # Find the generated pdf in the dir
            generated_pdf = [f for f in os.listdir(output_dir) if f.endswith(".pdf")][0]
            pdf_to_docx(os.path.join(output_dir, generated_pdf), output_path)
        else:
            raise Exception(f"Conversion failed: {str(e)}")

    return output_path

